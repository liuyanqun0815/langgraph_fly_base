"""Abstract interface for document loader implementations."""
import datetime
import mimetypes
import os
import tempfile
import uuid
from urllib.parse import urlparse

import requests
from docx import Document as DocxDocument
from langchain_core.documents import Document

from config import get_env
from sale_app.core.kb.loader.extractor_base import BaseExtractor

class WordExtractor(BaseExtractor):
    """
    用于加载 docx 文件的提取器。

    参数:
        file_path (str): 需要加载的文件路径。
    """

    def __init__(self, file_path: str):
        """
        使用指定的文件路径初始化 WordExtractor。

        参数:
            file_path (str): 文件路径。
        """
        self.file_path = file_path

        # 将路径中的“~”扩展为完整的用户家目录路径。
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # 如果给定的路径不是本地文件且是一个有效的URL，则下载文件到临时位置并使用该临时文件。
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

                    # 如果请求返回的状态码不是 200，则抛出带有错误信息的异常
        if r.status_code != 200:
            raise ValueError(
                f"请检查您的文件 URL；返回的状态码为 {r.status_code}"
            )

            # 当文件下载成功后，更新文件路径，创建一个临时文件并将响应的内容写入其中
            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
            # 如果提供的文件路径既不是有效的本地文件也不是有效的 URL，则抛出异常
        elif not os.path.isfile(self.file_path):
            raise ValueError(f"文件路径 {self.file_path} 不是有效的文件或 URL")

    # 析构方法，用于关闭临时文件
    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    # 加载指定路径作为单页内容
    def extract(self) -> list[Document]:
        """加载给定路径作为单页内容。

        Returns:
            list[Document]: 包含页面内容和元数据的文档列表。
        """
        content = self.parse_docx(self.file_path, 'storage')
        return [Document(
            page_content=content,
            metadata={"source": self.file_path},
        )]

    # 静态方法，用于检查 URL 是否有效
    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """检查 URL 是否有效。

        Args:
            url (str): 要检查的 URL。

        Returns:
            bool: 如果 URL 有效则返回 True，否则返回 False。
        """
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    # 从给定的 docx 文档中提取图片，并将这些图片保存到指定的文件夹中
    def _extract_images_from_docx(self, doc, image_folder):
        """
        从给定的 docx 文档中提取图片，并将这些图片保存到指定的文件夹中。

        参数：
            doc：docx 文档对象
            image_folder：存储图片的文件夹路径

        返回：
            一个字典，包含了 docx 文档中的所有图片。字典的键是图片在文档中的位置，值是图片的 URL。
        """
        os.makedirs(image_folder, exist_ok=True)
        image_count = 0
        image_map = {}

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                image_ext = rel.target_ref.split('.')[-1]
                # 使用 UUID 作为文件名
                file_uuid = str(uuid.uuid4())
                file_key = 'image_files/' + file_uuid + '.' + image_ext
                mime_type, _ = mimetypes.guess_type(file_key)

                image_map[
                    rel.target_part] = f"![image]({get_env('SERVICE_API_URL')}/files/image-preview/{file_uuid + '.' + image_ext})"

        return image_map

    # 将表格转换为 Markdown 格式
    def _table_to_markdown(self, table):
        markdown = ""
        # 处理表格头部
        header_row = table.rows[0]
        headers = [cell.text for cell in header_row.cells]
        markdown += "| " + " | ".join(headers) + " |\n"
        markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        # 处理表格行
        for row in table.rows[1:]:
            row_cells = [cell.text for cell in row.cells]
            markdown += "| " + " | ".join(row_cells) + " |\n"

        return markdown

    # 解析段落，处理段落中的文本和图片
    def _parse_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip'):
                for blip in run.element.xpath('.//a:blip'):
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return ' '.join(paragraph_content) if paragraph_content else ''

    # 解析 docx 文件
    def parse_docx(self, docx_path, image_folder):
        doc = DocxDocument(docx_path)
        os.makedirs(image_folder, exist_ok=True)

        content = []

        image_map = self._extract_images_from_docx(doc, image_folder)

        def parse_paragraph(paragraph):
            paragraph_content = []
            for run in paragraph.runs:
                if run.element.tag.endswith('r'):
                    drawing_elements = run.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        for blip in blip_elements:
                            embed_id = blip.get(
                                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    paragraph_content.append(image_map[image_part])
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return ''.join(paragraph_content) if paragraph_content else ''

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                para = paragraphs.pop(0)
                parsed_paragraph = parse_paragraph(para)
                if parsed_paragraph:
                    content.append(parsed_paragraph)
            elif element.tag.endswith('tbl'):  # 表格
                table = tables.pop(0)
                content.append(self._table_to_markdown(table))
        return '\n'.join(content)


